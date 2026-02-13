# docx_ops/replace_project.py
from __future__ import annotations
from typing import List
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def replace_first_project_safely(doc: Document, new_title: str, new_bullets: List[str]) -> Document:
    def delete_paragraph(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        paragraph._p = paragraph._element = None

    def format_title(paragraph, text):
        run = paragraph.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)

    def format_bullet(paragraph, text):
        run = paragraph.add_run(f"• {text}")
        run.font.size = Pt(10.5)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.15)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)

    def is_section_header(text: str) -> bool:
        t = (text or "").strip()
        # e.g., "WORK EXPERIENCE", "EDUCATION"
        return len(t) >= 4 and t.upper() == t and any(c.isalpha() for c in t)

    def is_bullet_line(text: str) -> bool:
        t = (text or "").strip()
        return t.startswith("•") or t.startswith("-")

    def is_title_like(paragraph) -> bool:
        # Common case: title is bolded somewhere in the paragraph
        if any(run.bold for run in paragraph.runs if run.text.strip()):
            return True
        # Backup heuristic: title lines often contain a date separator like "|" or "—"
        t = (paragraph.text or "").strip()
        if "|" in t or "–" in t or "-" in t:
            # avoid treating bullet lines as titles
            return not is_bullet_line(t)
        return False

    new_bullets = [bp.strip() for bp in new_bullets if bp and bp.strip()]

    # 1) Find PROJECT EXPERIENCE header
    section_found = False
    header_idx = -1
    for i, para in enumerate(doc.paragraphs):
        if "PROJECT EXPERIENCE" in (para.text or "").upper():
            section_found = True
            header_idx = i
            break

    if not section_found:
        raise ValueError("Could not find 'PROJECT EXPERIENCE' section in the document.")

    # 2) Find first non-empty paragraph after header => start of first project block
    start_idx = -1
    for i in range(header_idx + 1, len(doc.paragraphs)):
        if doc.paragraphs[i].text.strip():
            start_idx = i
            break

    if start_idx == -1:
        raise ValueError("Found 'PROJECT EXPERIENCE' but couldn't locate the first project entry below it.")

    # 3) Find end of first project block using robust heuristics
    end_idx = len(doc.paragraphs)  # default fallback
    seen_bullet = False

    for i in range(start_idx + 1, len(doc.paragraphs)):
        txt = doc.paragraphs[i].text.strip()

        if not txt:
            # blank line after we've seen bullets => likely end of block
            if seen_bullet:
                end_idx = i
                break
            continue

        if is_section_header(txt):
            end_idx = i
            break

        if is_bullet_line(txt):
            seen_bullet = True
            continue

        # A new "title-like" line after we've started (and especially after bullets) => next project
        if is_title_like(doc.paragraphs[i]) and (seen_bullet or i == start_idx + 1):
            end_idx = i
            break

    # 4) Delete old block
    for idx in reversed(range(start_idx, end_idx)):
        delete_paragraph(doc.paragraphs[idx])

    # 5) Insert new block at start_idx
    insert_anchor = doc.paragraphs[start_idx] if start_idx < len(doc.paragraphs) else doc.paragraphs[-1]

    # Insert title then bullets (in reverse using insert_paragraph_before)
    for bullet in reversed(new_bullets):
        p = insert_anchor.insert_paragraph_before("")
        format_bullet(p, bullet)

    p_title = insert_anchor.insert_paragraph_before("")
    format_title(p_title, new_title)

    return doc
