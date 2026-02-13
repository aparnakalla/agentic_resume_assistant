from pathlib import Path
import sys

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx_ops.replace_project import replace_first_project_safely


def test_replace_first_project_removes_summary_and_replaces_first_project_block():
    doc = Document()
    doc.add_paragraph("WORK EXPERIENCE")
    doc.add_paragraph("Prior role")

    doc.add_paragraph("PROJECT EXPERIENCE")

    doc.add_paragraph("Built internal tooling that accelerated onboarding.")
    doc.add_paragraph("Collaborated across teams to improve delivery quality.")
    doc.add_paragraph("Mentored junior engineers on code review best practices.")

    first_project_title = doc.add_paragraph("Legacy Project Title")
    first_project_title.runs[0].bold = True
    doc.add_paragraph("• Legacy first bullet")
    doc.add_paragraph("• Legacy second bullet")

    second_project_title = doc.add_paragraph("Second Project")
    second_project_title.runs[0].bold = True
    doc.add_paragraph("• Existing second-project bullet")

    replace_first_project_safely(
        doc,
        new_title="New Project Title",
        new_bullets=["Shipped feature X", "Reduced latency by 40%"],
    )

    texts = [p.text for p in doc.paragraphs]

    assert "Built internal tooling that accelerated onboarding." not in texts
    assert "Collaborated across teams to improve delivery quality." not in texts
    assert "Mentored junior engineers on code review best practices." not in texts
    assert "Legacy Project Title" not in texts
    assert "• Legacy first bullet" not in texts
    assert "• Legacy second bullet" not in texts

    project_header_idx = texts.index("PROJECT EXPERIENCE")
    second_project_idx = texts.index("Second Project")

    assert texts[project_header_idx + 1 : second_project_idx] == [
        "New Project Title",
        "• Shipped feature X",
        "• Reduced latency by 40%",
    ]
